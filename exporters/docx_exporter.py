import docx
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def generate_docx(resume_json: dict, output_path: str) -> str:
    """
    Generates a clean, ATS-friendly DOCX resume from the parsed JSON data.
    Safely handles both dictionary and string formats returned by LLMs.
    """
    doc = docx.Document()
    
    # 1. Header (Name and Contact Info)
    name = doc.add_heading(resume_json.get("name", "Name Not Provided"), 0)
    name.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    email = resume_json.get("email", "")
    phone = resume_json.get("phone", "")
    
    # Only add contact info if it exists
    contact_str = " | ".join(filter(None, [email, phone]))
    if contact_str:
        contact_info = doc.add_paragraph(contact_str)
        contact_info.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # 2. Skills Section
    skills_list = resume_json.get("skills", [])
    if skills_list:
        doc.add_heading("Skills", level=1)
        # Ensure skills is a list of strings
        if isinstance(skills_list, list):
            doc.add_paragraph(", ".join([str(s) for s in skills_list]))
        elif isinstance(skills_list, str):
            doc.add_paragraph(skills_list)
        
    # 3. Experience Section
    experience_list = resume_json.get("experience", [])
    if experience_list:
        doc.add_heading("Professional Experience", level=1)
        for exp in experience_list:
            if isinstance(exp, dict):
                p = doc.add_paragraph()
                p.add_run(f"{exp.get('role', 'Role')}").bold = True
                p.add_run(f" | {exp.get('company', 'Company')}").italic = True
                
                duration = exp.get("duration", "")
                if duration:
                    p.add_run(f" ({duration})")
                
                # Handle bullets safely
                bullets = exp.get("bullets", [])
                if isinstance(bullets, list):
                    for bullet in bullets:
                        doc.add_paragraph(str(bullet), style='List Bullet')
                elif isinstance(bullets, str):
                    doc.add_paragraph(bullets, style='List Bullet')
                    
            elif isinstance(exp, str):
                # If LLM returned a plain string instead of dict
                doc.add_paragraph(exp, style='List Bullet')
            
    # 4. Education Section
    education_list = resume_json.get("education", [])
    if education_list:
        doc.add_heading("Education", level=1)
        for edu in education_list:
            p = doc.add_paragraph()
            
            if isinstance(edu, dict):
                p.add_run(f"{edu.get('degree', 'Degree')}").bold = True
                
                institution = edu.get("institution", edu.get("school", ""))
                year = edu.get("year", edu.get("date", ""))
                
                edu_text = ""
                if institution:
                    edu_text += f" - {institution}"
                if year:
                    edu_text += f" ({year})"
                    
                if edu_text:
                    p.add_run(edu_text)
                    
            elif isinstance(edu, str):
                # If LLM returned a plain string
                p.add_run(edu).bold = True
        
    # 5. Projects Section
    projects_list = resume_json.get("projects", [])
    if projects_list:
        doc.add_heading("Projects", level=1)
        for proj in projects_list:
            if isinstance(proj, dict):
                p = doc.add_paragraph()
                p.add_run(f"{proj.get('name', 'Project Name')}").bold = True
                
                tech = proj.get("technologies", [])
                if isinstance(tech, list) and tech:
                    p.add_run(f" (Technologies: {', '.join([str(t) for t in tech])})").italic = True
                elif isinstance(tech, str) and tech:
                    p.add_run(f" (Technologies: {tech})").italic = True
                    
                desc = proj.get("description", "")
                if isinstance(desc, str) and desc:
                    doc.add_paragraph(desc)
                elif isinstance(desc, list): # Sometimes descriptions are bulleted lists
                    for d in desc:
                        doc.add_paragraph(str(d), style='List Bullet')
                        
            elif isinstance(proj, str):
                # If LLM returned a plain string
                doc.add_paragraph(proj, style='List Bullet')

    # Save the document
    doc.save(output_path)
    return output_path